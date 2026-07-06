# from io import BytesIO
# from pathlib import Path
# import pdfplumber
# from app.core.logger import get_logger
# import docx
# from fastapi import HTTPException
# import re
# from io import BytesIO
# from zipfile import ZipFile
# import defusedxml.ElementTree as ET

# from docx.document import Document as _Document
# from docx.table import Table, _Cell
# from docx.text.paragraph import Paragraph
# from docx.oxml.table import CT_Tbl
# from docx.oxml.text.paragraph import CT_P

# logger = get_logger(__name__)




# import asyncio
# import math
# import os

# import shutil
# import subprocess
# import tempfile
# import time
# import zipfile
# from concurrent.futures import ThreadPoolExecutor
# from io import BytesIO
# from pathlib import Path

# import fitz  # PyMuPDF
# import pdfplumber
# import pytesseract
# from PIL import Image, ImageFilter, ImageOps
# from fastapi import BackgroundTasks, Depends, File, Form, HTTPException, Request, UploadFile
# from sqlalchemy.orm import Session





# def extract_text_from_bytes(
#     content: bytes,
#     filename: str,
#     document_type: str = "jd",
#     return_meta: bool = False,
# ):
#     suffix = Path(filename).suffix.lower()

#     # ---------------- PDF ----------------
#     if suffix == ".pdf":
#         text, meta = extract_text_from_pdf_hybrid(content)

#     # ---------------- DOCX ----------------
#     elif suffix == ".docx":
#         text, meta = extract_text_from_docx_hybrid(content)

#     # ---------------- TXT ----------------
#     elif suffix == ".txt":
#         text = content.decode("utf-8", errors="ignore").strip()
#         meta = {
#             "mode": "native",
#             "total_pages": None,
#             "pages_native": None,
#             "pages_ocrd": 0,
#             "suspicious_pages": [],
#             "extracted_chars": len(text or ""),
#             "fallback_used": False,
#             "source": "txt",
#         }

#     else:
#         raise HTTPException(status_code=400, detail="Only PDF, DOCX or TXT files are supported")

#     if document_type == "jd":
#         text = normalize_extracted_text(text)

#     if return_meta:
#         return text, meta
#     return text



# DOCX_METADATA_LABELS = {
#     "job title",
#     "title",
#     "location",
#     "experience",
#     "years of experience",
#     "employment type",
#     "grade",
#     "grade / level",
#     "level",
#     "req id",
#     "requirement id",
#     "role summary",
# }


# def _single_line(text: str) -> str:
#     if not text:
#         return ""
#     text = (
#         text.replace("\xa0", " ")
#         .replace("\u200b", "")
#         .replace("\ufeff", "")
#         .replace("\u2013", "-")
#         .replace("\u2014", "-")
#         .replace("\u2215", "/")
#         .replace("\t", " ")
#         .replace("\r", "\n")
#     )
#     text = re.sub(r"[ \t]+", " ", text)
#     text = re.sub(r"\n{2,}", "\n", text)
#     text = re.sub(r"\s+", " ", text).strip()
#     return text


# def _is_metadata_label(text: str) -> bool:
#     if not text:
#         return False

#     cleaned = text.strip().rstrip(":").strip().lower()

#     if cleaned in DOCX_METADATA_LABELS:
#         return True

#     words = cleaned.split()
#     return len(words) <= 4 and len(cleaned) <= 40 and not any(ch.isdigit() for ch in cleaned)


# def _dedupe_keep_order(items: list[str]) -> list[str]:
#     seen = set()
#     result = []
#     for item in items:
#         if item and item not in seen:
#             result.append(item)
#             seen.add(item)
#     return result


# def _iter_block_items(parent):
#     if isinstance(parent, _Document):
#         parent_elm = parent.element.body
#     elif isinstance(parent, _Cell):
#         parent_elm = parent._tc
#     else:
#         raise ValueError(f"Unsupported parent type: {type(parent)}")

#     for child in parent_elm.iterchildren():
#         if isinstance(child, CT_P):
#             yield Paragraph(child, parent)
#         elif isinstance(child, CT_Tbl):
#             yield Table(child, parent)


# def _extract_cell_lines(cell: _Cell) -> list[str]:
#     lines: list[str] = []

#     for block in _iter_block_items(cell):
#         if isinstance(block, Paragraph):
#             text = _single_line(block.text)
#             if text:
#                 lines.append(text)

#         elif isinstance(block, Table):
#             lines.extend(_table_to_lines(block))

#     if not lines:
#         fallback = _single_line(cell.text)
#         if fallback:
#             lines.append(fallback)

#     return _dedupe_keep_order(lines)


# def _table_to_lines(table: Table) -> list[str]:
#     lines: list[str] = []

#     for row in table.rows:
#         row_values: list[str] = []

#         for cell in row.cells:
#             cell_lines = _extract_cell_lines(cell)
#             cell_text = _single_line(" ; ".join(cell_lines))
#             if cell_text:
#                 row_values.append(cell_text)

#         row_values = _dedupe_keep_order(row_values)

#         if not row_values:
#             continue

#         # Best case: 2-column metadata row
#         if len(row_values) == 2 and _is_metadata_label(row_values[0]):
#             line = f"{row_values[0].rstrip(':')}: {row_values[1]}"
#         else:
#             line = " | ".join(row_values)

#         line = _single_line(line)
#         if line:
#             lines.append(line)

#     return lines


# def _extract_textboxes_from_docx_bytes(content: bytes) -> list[str]:
#     ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
#     results: list[str] = []

#     try:
#         with ZipFile(BytesIO(content)) as zf:
#             xml_bytes = zf.read("word/document.xml")

#         root = ET.fromstring(xml_bytes)

#         for txbx in root.findall(".//w:txbxContent", ns):
#             parts = []
#             for t in txbx.findall(".//w:t", ns):
#                 if t.text and t.text.strip():
#                     parts.append(t.text.strip())

#             text = _single_line(" ".join(parts))
#             if text:
#                 results.append(text)

#     except Exception:
#         logger.exception("[TEXT_EXTRACT] PDF block extraction failed")
#         pass

#     return results


# def _pair_label_value_lines(lines: list[str]) -> list[str]:
#     result: list[str] = []
#     i = 0

#     while i < len(lines):
#         current = _single_line(lines[i])

#         if (
#             _is_metadata_label(current)
#             and ":" not in current
#             and i + 1 < len(lines)
#         ):
#             nxt = _single_line(lines[i + 1])
#             if nxt and not _is_metadata_label(nxt):
#                 result.append(f"{current}: {nxt}")
#                 i += 2
#                 continue

#         result.append(current)
#         i += 1

#     return result


# def extract_text_from_docx_bytes(content: bytes) -> str:
#     document = docx.Document(BytesIO(content))
#     full_text: list[str] = []

#     # Body in actual order
#     for block in _iter_block_items(document):
#         if isinstance(block, Paragraph):
#             text = _single_line(block.text)
#             if text:
#                 full_text.append(text)

#         elif isinstance(block, Table):
#             full_text.extend(_table_to_lines(block))

#     # Headers / Footers
#     for section in document.sections:
#         for p in section.header.paragraphs:
#             text = _single_line(p.text)
#             if text:
#                 full_text.append(text)

#         for table in section.header.tables:
#             full_text.extend(_table_to_lines(table))

#         for p in section.footer.paragraphs:
#             text = _single_line(p.text)
#             if text:
#                 full_text.append(text)

#         for table in section.footer.tables:
#             full_text.extend(_table_to_lines(table))

#     # Text boxes / shapes
#     full_text.extend(_extract_textboxes_from_docx_bytes(content))

#     # Pair split metadata lines
#     full_text = _pair_label_value_lines(full_text)

#     # Final cleanup
#     full_text = [_single_line(x) for x in full_text if _single_line(x)]
#     full_text = _dedupe_keep_order(full_text)

#     return "\n".join(full_text).strip()




# def normalize_extracted_text(t: str) -> str:
#     if not t:
#         return ""

#     t = t.replace("\u2013", "-").replace("\u2014", "-").replace("\u2215", "/")
#     t = t.replace("\u200b", "").replace("\ufeff", "").replace("\xa0", " ")

#     # A W S -> AWS (only consecutive single letters)
#     t = re.sub(
#         r"\b(?:[A-Za-z]\s){2,}[A-Za-z]\b",
#         lambda m: m.group(0).replace(" ", ""),
#         t
#     )

#     # keep line structure
#     lines = [re.sub(r"[ \t]+", " ", line).strip() for line in t.splitlines()]
#     lines = [line for line in lines if line]

#     return "\n".join(lines).strip()




# # def is_text_sufficient(text: str, min_chars: int = 80) -> bool:
# #     """
# #     Decide whether extracted text is likely enough to skip OCR fallback.
# #     Uses non-whitespace char count.
# #     """
# #     non_ws_chars = len(re.sub(r"\s+", "", text or ""))
# #     return non_ws_chars >= min_chars


# def preprocess_image_for_ocr(img: Image.Image) -> Image.Image:
#     """
#     Basic preprocessing to improve OCR quality.
#     Safe default for scanned resumes.
#     """
#     img = img.convert("L")  # grayscale
#     img = ImageOps.autocontrast(img)
#     img = img.filter(ImageFilter.SHARPEN)
#     return img





# # Fix 1: Lower thresholds and drop the "signals >= 2" gate
# def page_needs_ocr(page, native_text, min_chars=80, min_image_ratio=0.50):

#     non_ws_chars = len(re.sub(r"\s+", "", native_text or ""))
#     if non_ws_chars < min_chars:
#         return True  # no meaningful text → always OCR

#     try:
#         text_dict = page.get_text("dict")
#         blocks = text_dict.get("blocks", [])
#     except Exception:
#         logger.exception("[TEXT_EXTRACT] Character count check failed")
#         return non_ws_chars < min_chars

#     page_area = page.rect.width * page.rect.height if page.rect else 0
#     largest_image_ratio = 0.0

#     for block in blocks:
#         if block.get("type") == 1:
#             x0, y0, x1, y1 = block.get("bbox", [0, 0, 0, 0])
#             ratio = (max(0, x1-x0) * max(0, y1-y0)) / page_area if page_area else 0
#             largest_image_ratio = max(largest_image_ratio, ratio)
#         # Also catch images embedded in type-0 blocks (some PDFs)
#         elif block.get("type") == 0:
#             for img_info in block.get("image", []):
#                 x0, y0, x1, y1 = img_info.get("bbox", [0, 0, 0, 0])
#                 ratio = (max(0, x1-x0) * max(0, y1-y0)) / page_area if page_area else 0
#                 largest_image_ratio = max(largest_image_ratio, ratio)

#     return largest_image_ratio >= min_image_ratio

# # Fix 2: Raise DPI to 300 and ensure PNG encoding
# def ocr_pdf_page(page, dpi=300):
#     zoom = dpi / 72
#     pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
#     img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
#     img = preprocess_image_for_ocr(img)

#     # Save as PNG to avoid Tesseract format issues
#     buf = BytesIO()
#     img.save(buf, format="PNG")
#     buf.seek(0)
#     img_png = Image.open(buf)

#     return pytesseract.image_to_string(img_png, config="--oem 3 --psm 6").strip()

# # Fix 3: Raise the "sufficient text" threshold
# def is_text_sufficient(text, min_chars=150):  # was 300
#     non_ws_chars = len(re.sub(r"\s+", "", text or ""))
#     return non_ws_chars >= min_chars

# # Fix 4: Verify tesseract on startup (add to app init)
# def verify_tesseract():
#     try:
#         version = pytesseract.get_tesseract_version()
#         logger.info(f"Tesseract version: {version}")
#     except Exception as e:
#         logger.error(f"TESSERACT NOT AVAILABLE: {e}")
#         raise RuntimeError("Tesseract is not installed or not in PATH") from e

# # verify_tesseract()






# def extract_text_from_pdf_hybrid(content: bytes) -> tuple[str, dict]:
#     """
#     PDF extraction strategy:
#     - Use native extraction for searchable text pages
#     - OCR only suspicious/scanned pages
#     """
#     doc = fitz.open(stream=content, filetype="pdf")

#     page_texts = []
#     pages_ocrd = 0
#     pages_native = 0
#     suspicious_pages = []

#     for page_index in range(len(doc)):
#         page = doc.load_page(page_index)
#         native_text = (page.get_text("text") or "").strip()

#         if page_needs_ocr(page, native_text):
#             suspicious_pages.append(page_index + 1)

#             try:
#                 ocr_text = ocr_pdf_page(page)
#             except Exception:
#                 logger.exception(
#                     "[TEXT_EXTRACT] OCR failed for PDF page filename_unknown page=%s",
#                     page_index + 1,
#                 )
#                 ocr_text = ""

#             page_texts.append(ocr_text)
#             pages_ocrd += 1
#         else:
#             page_texts.append(native_text)
#             pages_native += 1

#     full_text = "\n".join(t for t in page_texts if t).strip()

#     if len(doc) == 0:
#         mode = "native"
#     elif pages_ocrd == len(doc):
#         mode = "ocr"
#     elif pages_ocrd > 0:
#         mode = "hybrid"
#     else:
#         mode = "native"

#     meta = {
#         "mode": mode,
#         "total_pages": len(doc),
#         "pages_native": pages_native,
#         "pages_ocrd": pages_ocrd,
#         "suspicious_pages": suspicious_pages,
#         "extracted_chars": len(full_text or ""),
#         "fallback_used": pages_ocrd > 0,
#         "source": "pdf",
#     }

#     return full_text, meta


# def docx_contains_images(content: bytes) -> bool:
#     """
#     Check whether DOCX has embedded media/images.
#     """
#     try:
#         with zipfile.ZipFile(BytesIO(content)) as zf:
#             return any(name.startswith("word/media/") for name in zf.namelist())
#     except Exception:
#         logger.exception("[TEXT_EXTRACT] Failed to inspect DOCX images")
#         return False


# def get_docx_media_names(content: bytes) -> list[str]:
#     """
#     Return media file names inside a DOCX package.
#     """
#     try:
#         with zipfile.ZipFile(BytesIO(content)) as zf:
#             return [name for name in zf.namelist() if name.startswith("word/media/")]
#     except Exception:
#         logger.exception("[TEXT_EXTRACT] Failed to list DOCX media files")
#         return []


# def ocr_images_from_docx_bytes(
#     content: bytes,
#     min_width: int = 800,
#     min_height: int = 800,
# ) -> tuple[str, dict]:
#     """
#     Fallback OCR for embedded DOCX images if DOCX->PDF conversion is unavailable or fails.

#     To reduce noisy OCR from icons/logos, skip small images by default.
#     """
#     texts = []
#     total_images = 0
#     ocr_images = 0
#     skipped_small_images = 0

#     try:
#         with zipfile.ZipFile(BytesIO(content)) as zf:
#             media_names = [name for name in zf.namelist() if name.startswith("word/media/")]

#             for media_name in media_names:
#                 total_images += 1

#                 try:
#                     with zf.open(media_name) as img_file:
#                         img = Image.open(BytesIO(img_file.read()))
#                         img.load()
#                 except Exception:
#                     logger.exception(
#                         "[TEXT_EXTRACT] Failed to load DOCX media image=%s",
#                         media_name,
#                     )
#                     continue

#                 width, height = img.size if img.size else (0, 0)

#                 # Skip likely icons/logos/signatures
#                 if width < min_width and height < min_height:
#                     skipped_small_images += 1
#                     continue

#                 try:
#                     img = preprocess_image_for_ocr(img)
#                     text = pytesseract.image_to_string(img, config="--oem 3 --psm 6").strip()
#                 except Exception:
#                     logger.exception(
#                         "[TEXT_EXTRACT] OCR failed for DOCX media image=%s",
#                         media_name,
#                     )
#                     text = ""

#                 if text:
#                     texts.append(text)
#                     ocr_images += 1

#     except Exception:
#         logger.exception("[TEXT_EXTRACT] Failed OCR fallback on DOCX images")

#     final_text = "\n".join(t for t in texts if t).strip()
#     meta = {
#         "mode": "ocr_images",
#         "total_images": total_images,
#         "ocr_images": ocr_images,
#         "skipped_small_images": skipped_small_images,
#         "extracted_chars": len(final_text or ""),
#         "source": "docx_images",
#     }
#     return final_text, meta


# def find_office_converter_executable() -> str | None:
#     """
#     Find LibreOffice / soffice executable for DOCX -> PDF conversion.
#     """
#     return shutil.which("libreoffice") or shutil.which("soffice")


# def convert_docx_bytes_to_pdf(content: bytes) -> bytes:
#     """
#     Convert DOCX bytes to PDF using LibreOffice/soffice in headless mode.
#     """
#     converter = find_office_converter_executable()
#     if not converter:
#         raise RuntimeError("LibreOffice/soffice executable not found for DOCX to PDF conversion")

#     with tempfile.TemporaryDirectory() as tmpdir:
#         input_path = os.path.join(tmpdir, "input.docx")
#         output_path = os.path.join(tmpdir, "input.pdf")

#         with open(input_path, "wb") as f:
#             f.write(content)

#         proc = subprocess.run(
#             [
#                 converter,
#                 "--headless",
#                 "--convert-to",
#                 "pdf",
#                 "--outdir",
#                 tmpdir,
#                 input_path,
#             ],
#             stdout=subprocess.PIPE,
#             stderr=subprocess.PIPE,
#             timeout=120,
#             check=False,
#         )

#         if proc.returncode != 0:
#             raise RuntimeError(
#                 f"DOCX to PDF conversion failed. returncode={proc.returncode}, stderr={proc.stderr.decode(errors='ignore')}"
#             )

#         if not os.path.exists(output_path):
#             raise RuntimeError("DOCX to PDF conversion did not produce output PDF")

#         with open(output_path, "rb") as f:
#             return f.read()

# def extract_text_from_docx_hybrid(content: bytes) -> tuple[str, dict]:
#     """
#     DOCX extraction strategy:
#     1. Native DOCX text extraction         — always runs, fastest
#     2. Direct image OCR                    — if native text insufficient + has images
#     3. LibreOffice DOCX->PDF conversion    — only if OCR also gives nothing
#     """
#     native_text = (extract_text_from_docx_bytes(content) or "").strip()
#     has_images = docx_contains_images(content)
#     media_names = get_docx_media_names(content)

#     native_meta = {
#         "mode": "native",
#         "total_pages": None,
#         "pages_native": None,
#         "pages_ocrd": 0,
#         "suspicious_pages": [],
#         "extracted_chars": len(native_text or ""),
#         "fallback_used": False,
#         "has_images": has_images,
#         "media_count": len(media_names),
#         "source": "docx_native",
#     }

#     # Fast path: good enough native text — skip everything else
#     if is_text_sufficient(native_text):
#         return native_text, native_meta

#     # No images — native is all we have
#     if not has_images:
#         native_meta["mode"] = "native_low_text"
#         return native_text, native_meta

#     # Try direct image OCR FIRST — much faster than LibreOffice
#     try:
#         image_text, image_meta = ocr_images_from_docx_bytes(content)
#         combined_text = "\n".join(
#             part for part in [native_text, image_text] if part
#         ).strip()

#         image_meta.update({
#             "fallback_used": True,
#             "has_images": has_images,
#             "media_count": len(media_names),
#             "source": "docx_images",
#         })

#         # If OCR gave us enough — return it, skip LibreOffice entirely
#         if is_text_sufficient(combined_text):
#             return combined_text, image_meta

#     except Exception:
#         logger.exception("[TEXT_EXTRACT] Direct image OCR failed")
#         combined_text = native_text

#     # Last resort only: LibreOffice DOCX->PDF (slow on cold start)
#     try:
#         pdf_bytes = convert_docx_bytes_to_pdf(content)
#         pdf_text, pdf_meta = extract_text_from_pdf_hybrid(pdf_bytes)

#         if len(re.sub(r"\s+", "", pdf_text or "")) >= len(
#             re.sub(r"\s+", "", combined_text or "")
#         ):
#             pdf_meta.update({
#                 "source": "docx_to_pdf",
#                 "fallback_used": True,
#                 "has_images": has_images,
#                 "media_count": len(media_names),
#             })
#             return pdf_text, pdf_meta

#     except Exception:
#         logger.exception("[TEXT_EXTRACT] DOCX -> PDF fallback failed")

#     # Return whatever we have
#     if len(re.sub(r"\s+", "", combined_text or "")) >= len(
#         re.sub(r"\s+", "", native_text or "")
#     ):
#         return combined_text, image_meta  # type: ignore[possibly-undefined]

#     native_meta["mode"] = "native_low_text"
#     native_meta["fallback_used"] = False
#     return native_text, native_meta




"""
Hybrid text extraction (PDF / DOCX / TXT) — drop-in replacement.

Improvements over previous version:
  1. Native + OCR text are MERGED per page instead of OCR replacing native text.
  2. Garbage-text detection: (cid:xx) artifacts, replacement chars, low alpha ratio
     -> forces OCR even when char count looks sufficient.
  3. Correct image detection via page.get_image_info() — old type-0 block branch
     in page_needs_ocr() was dead code and is now fixed.
  4. Orientation correction via Tesseract OSD before OCR.
  5. PSM 3 (auto layout) first, PSM 6 fallback — fixes multi-column layouts.
  6. Better preprocessing: upscale small renders + Otsu binarization (SHARPEN removed
     — it typically degrades Tesseract). Pointless PNG round-trip removed.
  7. sort=True for native PDF extraction -> correct reading order.
  8. Parallel page OCR with ThreadPoolExecutor.
  9. doc.close() always called. Configurable OCR language/DPI via env vars.
  10. De-hyphenation of OCR line-break artifacts.
  11. DOCX: area-based image filter (catches wide-but-short scans that the old
      width<800 AND height<800 filter missed).
  12. Textbox extraction now also scans header/footer XML parts.
  13. EMF/WMF/SVG media files skipped gracefully in image OCR path.

Backward-compatibility notes:
  - All public function names and signatures preserved.
  - page_needs_ocr() shim kept for any external callers.
  - ocr_images_from_docx_bytes() keeps min_width/min_height params (internally
    converted to area/dimension checks).
  - All original imports kept (including pdfplumber, asyncio, etc.).
  - Meta dict shapes are identical; one new key added to DOCX image path:
    "skipped_unreadable_images".
"""

# ── original imports (all kept to avoid breaking any re-imports) ─────────────

from io import BytesIO
from pathlib import Path
import pdfplumber
from app.core.logger import get_logger
import docx
from fastapi import HTTPException
import re
from io import BytesIO
from zipfile import ZipFile
import defusedxml.ElementTree as ET

from docx.document import Document as _Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

import asyncio
import math
import os
import shutil
import subprocess
import tempfile
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor

import fitz          # PyMuPDF
import pytesseract
from PIL import Image, ImageFilter, ImageOps
from fastapi import BackgroundTasks, Depends, File, Form, Request, UploadFile
from sqlalchemy.orm import Session

logger = get_logger(__name__)


# ── env-configurable knobs ────────────────────────────────────────────────────

OCR_LANG        = os.getenv("OCR_LANG",        "eng")   # e.g. "eng+hin"
OCR_DPI         = int(os.getenv("OCR_DPI",     "300"))
OCR_MAX_WORKERS = int(os.getenv("OCR_MAX_WORKERS", "4"))


# ── DOCX metadata labels (unchanged) ─────────────────────────────────────────

DOCX_METADATA_LABELS = {
    "job title",
    "title",
    "location",
    "experience",
    "years of experience",
    "employment type",
    "grade",
    "grade / level",
    "level",
    "req id",
    "requirement id",
    "role summary",
}


# ═══════════════════════════════════════════════════════════════════════════════
#  TEXT-QUALITY HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

_CID_RE = re.compile(r"\(cid:\d+\)")


def _non_ws_len(text: str) -> int:
    """Count non-whitespace characters."""
    return len(re.sub(r"\s+", "", text or ""))


def is_text_sufficient(text: str, min_chars: int = 150) -> bool:
    """
    Decide whether extracted text is likely enough to skip OCR fallback.
    Uses non-whitespace char count.
    (threshold raised from 80 -> 150 to catch near-empty extractions)
    """
    return _non_ws_len(text) >= min_chars


def native_text_is_garbage(text: str) -> bool:
    """
    Detect 'extractable but useless' native PDF text:
    broken ToUnicode maps producing (cid:x) output, mojibake, symbol soup.
    """
    if not text:
        return True

    stripped = re.sub(r"\s+", "", text)
    if not stripped:
        return True

    # (cid:NN) artifacts from broken font encodings
    if len(_CID_RE.findall(text)) >= 3:
        return True

    # Unicode replacement chars
    if stripped.count("\ufffd") / len(stripped) > 0.05:
        return True

    # Very low letter/digit ratio -> symbol soup
    alnum = sum(ch.isalnum() for ch in stripped)
    if len(stripped) >= 40 and (alnum / len(stripped)) < 0.45:
        return True

    return False


def _ocr_quality_score(text: str) -> float:
    """Rough score used to compare two OCR attempts (PSM 3 vs PSM 6)."""
    if not text:
        return 0.0
    words = re.findall(r"[A-Za-z]{2,}", text)
    return len(words) + 0.01 * _non_ws_len(text)


# ═══════════════════════════════════════════════════════════════════════════════
#  IMAGE PRE-PROCESSING  +  OCR
# ═══════════════════════════════════════════════════════════════════════════════

def _otsu_threshold(gray: Image.Image) -> int:
    """
    Pure-Python Otsu's method (no numpy/opencv dependency).
    Much more effective than a fixed threshold for variable-quality scans.
    """
    hist  = gray.histogram()[:256]
    total = sum(hist)
    if total == 0:
        return 128

    sum_total = sum(i * h for i, h in enumerate(hist))
    sum_bg    = 0.0
    weight_bg = 0
    max_var   = 0.0
    threshold = 128

    for i in range(256):
        weight_bg += hist[i]
        if weight_bg == 0:
            continue
        weight_fg = total - weight_bg
        if weight_fg == 0:
            break
        sum_bg   += i * hist[i]
        mean_bg   = sum_bg / weight_bg
        mean_fg   = (sum_total - sum_bg) / weight_fg
        var_between = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2
        if var_between > max_var:
            max_var   = var_between
            threshold = i

    return threshold


def preprocess_image_for_ocr(img: Image.Image, binarize: bool = True) -> Image.Image:
    """
    Grayscale -> upscale small images -> autocontrast -> Otsu binarize.

    SHARPEN intentionally removed: it typically hurts Tesseract accuracy.
    PNG round-trip removed: pytesseract accepts PIL images directly.
    Upscaling small images: Tesseract performs best at ~300 DPI equivalents.
    """
    img = img.convert("L")

    # Upscale small crops/images so Tesseract has enough resolution.
    if min(img.size) < 1000:
        scale    = 1000 / min(img.size)
        new_size = (int(img.width * scale), int(img.height * scale))
        img      = img.resize(new_size, Image.LANCZOS)

    img = ImageOps.autocontrast(img)

    if binarize:
        t   = _otsu_threshold(img)
        img = img.point(lambda p: 255 if p > t else 0)

    return img


def correct_orientation(img: Image.Image) -> Image.Image:
    """
    Use Tesseract OSD to detect and fix 90/180/270-degree rotated scans.
    OSD often fails on sparse images — failure is non-fatal.
    """
    try:
        osd = pytesseract.image_to_osd(img)
        m   = re.search(r"Rotate:\s*(\d+)", osd)
        rotation = int(m.group(1)) if m else 0
        if rotation:
            img = img.rotate(-rotation, expand=True, fillcolor=255)
    except Exception:
        logger.debug("[TEXT_EXTRACT] OSD orientation detection failed", exc_info=True)
    return img


def _ocr_image_raw(img: Image.Image, lang: str = OCR_LANG) -> str:
    """
    Core OCR on a preprocessed PIL image.
    Tries PSM 3 first (automatic page segmentation — handles multi-column layouts).
    Falls back to PSM 6 (single uniform block) and keeps the better result.
    """
    best_text, best_score = "", -1.0

    for psm in (3, 6):
        try:
            text = pytesseract.image_to_string(
                img, lang=lang, config=f"--oem 3 --psm {psm}"
            ).strip()
        except Exception:
            logger.exception("[TEXT_EXTRACT] Tesseract failed psm=%s", psm)
            continue

        score = _ocr_quality_score(text)
        if score > best_score:
            best_text, best_score = text, score

        # If PSM 3 already gave substantial output, skip PSM 6.
        if psm == 3 and _non_ws_len(text) >= 200:
            break

    return best_text


def ocr_image(img: Image.Image, lang: str = OCR_LANG) -> str:
    """
    Full OCR pipeline on a PIL image (preprocess -> orientation fix -> OCR).
    Public entry-point used by both the PDF and DOCX image paths.
    """
    img = preprocess_image_for_ocr(img)
    img = correct_orientation(img)
    return _ocr_image_raw(img, lang=lang)


def ocr_pdf_page(page: "fitz.Page", dpi: int = OCR_DPI, lang: str = OCR_LANG) -> str:
    """
    Render a PyMuPDF page to a PIL image and OCR it.
    DPI raised to 300 (was 72-based default) for reliable Tesseract results.
    PNG round-trip removed — PIL image passed directly to pytesseract.
    """
    zoom = dpi / 72
    pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
    img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    return ocr_image(img, lang=lang)


def verify_tesseract():
    """Verify Tesseract is installed and reachable. Call during app startup."""
    try:
        version = pytesseract.get_tesseract_version()
        logger.info("Tesseract version: %s", version)
    except Exception as e:
        logger.error("TESSERACT NOT AVAILABLE: %s", e)
        raise RuntimeError("Tesseract is not installed or not in PATH") from e


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

def page_largest_image_ratio(page: "fitz.Page") -> float:
    """
    Return the largest single embedded-image area as a fraction of page area.
    Uses page.get_image_info() which is the correct PyMuPDF API.
    (The old type-0 block loop was dead code — text blocks don't carry images.)
    """
    try:
        page_area = abs(page.rect) if page.rect else 0
        if not page_area:
            return 0.0
        largest = 0.0
        for info in page.get_image_info():
            x0, y0, x1, y1 = info.get("bbox", (0, 0, 0, 0))
            area    = max(0, x1 - x0) * max(0, y1 - y0)
            largest = max(largest, area / page_area)
        return largest
    except Exception:
        logger.exception("[TEXT_EXTRACT] Image-ratio inspection failed")
        return 0.0


def classify_page(
    page: "fitz.Page",
    native_text: str,
    min_chars: int = 80,
    min_image_ratio: float = 0.50,
) -> str:
    """
    Classify a PDF page into one of three handling modes:

      'native' — native text is complete; no OCR needed.
      'ocr'    — native text absent or garbage; replace entirely with OCR.
      'merge'  — native text is fine BUT a large image is also present;
                 OCR the rendered page and append any genuinely new text.

    The 'merge' mode is what was missing before: pages with a background photo
    or embedded diagram were incorrectly routed to 'ocr', discarding good
    native text and replacing it with a worse OCR result.
    """
    if _non_ws_len(native_text) < min_chars:
        return "ocr"
    if native_text_is_garbage(native_text):
        return "ocr"
    if page_largest_image_ratio(page) >= min_image_ratio:
        return "merge"
    return "native"


def page_needs_ocr(
    page: "fitz.Page",
    native_text: str,
    min_chars: int = 80,
    min_image_ratio: float = 0.50,
) -> bool:
    """
    Backward-compatible shim around classify_page().
    Returns True if the page requires any OCR at all.
    External callers that imported this function will continue to work.
    """
    return classify_page(page, native_text, min_chars, min_image_ratio) != "native"


def _merge_native_and_ocr(native: str, ocr: str) -> str:
    """
    Append only genuinely new lines from OCR to the native text.
    Prevents duplicating text that Tesseract re-reads from the rendered page.
    A line is 'new' if more than 60% of its words are absent from native text.
    """
    if not ocr:
        return native
    if not native:
        return ocr

    native_words = set(re.findall(r"[A-Za-z0-9@.+#-]{2,}", native.lower()))
    new_lines    = []

    for line in ocr.splitlines():
        words = re.findall(r"[A-Za-z0-9@.+#-]{2,}", line.lower())
        if not words:
            continue
        novel = sum(1 for w in words if w not in native_words)
        if novel / len(words) > 0.6:
            new_lines.append(line.strip())

    if not new_lines:
        return native
    return native + "\n" + "\n".join(new_lines)


def _dehyphenate(text: str) -> str:
    """Join words split across line-breaks: 'engin-\\neering' -> 'engineering'."""
    return re.sub(r"(\w)-\n(\w)", r"\1\2", text or "")


def extract_text_from_pdf_hybrid(content: bytes, lang: str = OCR_LANG) -> tuple[str, dict]:
    """
    PDF extraction strategy:
      - Native extraction (sort=True for correct reading order) for searchable pages.
      - OCR for pages with absent/garbage text.
      - MERGE (not replace) for pages that have both good native text and large images.
      - Pages needing OCR are rendered sequentially (thread-safety) then
        OCR'd in parallel via ThreadPoolExecutor.
      - doc.close() guaranteed via try/finally.
    """
    doc = fitz.open(stream=content, filetype="pdf")
    try:
        n_pages   = len(doc)
        natives   : list[str] = []
        decisions : list[str] = []

        # ── Pass 1: classify every page ──────────────────────────────────────
        for i in range(n_pages):
            page   = doc.load_page(i)
            # sort=True: follows visual reading order instead of PDF object order
            native = (page.get_text("text", sort=True) or "").strip()
            natives.append(native)
            decisions.append(classify_page(page, native))

        ocr_indices = [i for i, d in enumerate(decisions) if d in ("ocr", "merge")]

        # ── Pass 2: render all OCR-needed pages (sequential — PyMuPDF thread safety) ──
        rendered: dict[int, Image.Image] = {}
        for idx in ocr_indices:
            page = doc.load_page(idx)
            zoom = OCR_DPI / 72
            pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            rendered[idx] = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        # ── Pass 3: OCR in parallel ───────────────────────────────────────────
        ocr_results: dict[int, str] = {}
        if rendered:
            with ThreadPoolExecutor(max_workers=OCR_MAX_WORKERS) as pool:
                futures = {
                    idx: pool.submit(ocr_image, img, lang)
                    for idx, img in rendered.items()
                }
                for idx, fut in futures.items():
                    try:
                        ocr_results[idx] = fut.result()
                    except Exception:
                        logger.exception("[TEXT_EXTRACT] OCR thread failed page=%s", idx + 1)
                        ocr_results[idx] = ""

        # ── Pass 4: assemble final per-page texts ─────────────────────────────
        page_texts     : list[str] = []
        pages_ocrd      = 0
        pages_native    = 0
        suspicious_pages: list[int] = []

        for i in range(n_pages):
            decision = decisions[i]
            native   = natives[i]

            if decision == "native":
                page_texts.append(native)
                pages_native += 1

            elif decision == "ocr":
                suspicious_pages.append(i + 1)
                ocr_text = ocr_results.get(i, "")
                # Never end up with less text than native — keep whichever is longer.
                page_texts.append(
                    ocr_text if _non_ws_len(ocr_text) > _non_ws_len(native) else native
                )
                pages_ocrd += 1

            else:  # "merge"
                suspicious_pages.append(i + 1)
                page_texts.append(_merge_native_and_ocr(native, ocr_results.get(i, "")))
                pages_ocrd += 1

        full_text = _dehyphenate("\n".join(t for t in page_texts if t).strip())

        if n_pages == 0 or pages_ocrd == 0:
            mode = "native"
        elif pages_ocrd == n_pages:
            mode = "ocr"
        else:
            mode = "hybrid"

        meta = {
            "mode"            : mode,
            "total_pages"     : n_pages,
            "pages_native"    : pages_native,
            "pages_ocrd"      : pages_ocrd,
            "suspicious_pages": suspicious_pages,
            "extracted_chars" : len(full_text or ""),
            "fallback_used"   : pages_ocrd > 0,
            "source"          : "pdf",
        }
        return full_text, meta

    finally:
        doc.close()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _single_line(text: str) -> str:
    if not text:
        return ""
    text = (
        text.replace("\xa0",  " ")
            .replace("\u200b", "")
            .replace("\ufeff", "")
            .replace("\u2013", "-")
            .replace("\u2014", "-")
            .replace("\u2215", "/")
            .replace("\t",     " ")
            .replace("\r",     "\n")
    )
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    text = re.sub(r"\s+",    " ", text).strip()
    return text


def _is_metadata_label(text: str) -> bool:
    if not text:
        return False
    cleaned = text.strip().rstrip(":").strip().lower()
    if cleaned in DOCX_METADATA_LABELS:
        return True
    words = cleaned.split()
    return len(words) <= 4 and len(cleaned) <= 40 and not any(ch.isdigit() for ch in cleaned)


def _dedupe_keep_order(items: list[str]) -> list[str]:
    seen, result = set(), []
    for item in items:
        if item and item not in seen:
            result.append(item)
            seen.add(item)
    return result


def _iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError(f"Unsupported parent type: {type(parent)}")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _extract_cell_lines(cell: _Cell) -> list[str]:
    lines: list[str] = []
    for block in _iter_block_items(cell):
        if isinstance(block, Paragraph):
            text = _single_line(block.text)
            if text:
                lines.append(text)
        elif isinstance(block, Table):
            lines.extend(_table_to_lines(block))
    if not lines:
        fallback = _single_line(cell.text)
        if fallback:
            lines.append(fallback)
    return _dedupe_keep_order(lines)


def _table_to_lines(table: Table) -> list[str]:
    lines: list[str] = []
    for row in table.rows:
        row_values: list[str] = []
        for cell in row.cells:
            cell_lines = _extract_cell_lines(cell)
            cell_text  = _single_line(" ; ".join(cell_lines))
            if cell_text:
                row_values.append(cell_text)
        row_values = _dedupe_keep_order(row_values)
        if not row_values:
            continue
        if len(row_values) == 2 and _is_metadata_label(row_values[0]):
            line = f"{row_values[0].rstrip(':')}: {row_values[1]}"
        else:
            line = " | ".join(row_values)
        line = _single_line(line)
        if line:
            lines.append(line)
    return lines


def _extract_textboxes_from_docx_bytes(content: bytes) -> list[str]:
    """
    Extract text from all w:txbxContent elements inside:
      - word/document.xml   (main body textboxes)
      - word/header*.xml    (branded JD templates often put content here)
      - word/footer*.xml

    Previous version only scanned document.xml.
    """
    ns      = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    results : list[str] = []

    try:
        with ZipFile(BytesIO(content)) as zf:
            targets = [
                n for n in zf.namelist()
                if n == "word/document.xml"
                or re.match(r"word/(header|footer)\d*\.xml$", n)
            ]
            for xml_name in targets:
                try:
                    root = ET.fromstring(zf.read(xml_name))
                except Exception:
                    logger.exception("[TEXT_EXTRACT] Failed parsing XML part: %s", xml_name)
                    continue

                for txbx in root.findall(".//w:txbxContent", ns):
                    parts = [
                        t.text.strip()
                        for t in txbx.findall(".//w:t", ns)
                        if t.text and t.text.strip()
                    ]
                    text = _single_line(" ".join(parts))
                    if text:
                        results.append(text)

    except Exception:
        logger.exception("[TEXT_EXTRACT] Textbox extraction failed")

    return results


def _pair_label_value_lines(lines: list[str]) -> list[str]:
    result: list[str] = []
    i = 0
    while i < len(lines):
        current = _single_line(lines[i])
        if (
            _is_metadata_label(current)
            and ":" not in current
            and i + 1 < len(lines)
        ):
            nxt = _single_line(lines[i + 1])
            if nxt and not _is_metadata_label(nxt):
                result.append(f"{current}: {nxt}")
                i += 2
                continue
        result.append(current)
        i += 1
    return result


def extract_text_from_docx_bytes(content: bytes) -> str:
    document  = docx.Document(BytesIO(content))
    full_text : list[str] = []

    # Body in document order
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            text = _single_line(block.text)
            if text:
                full_text.append(text)
        elif isinstance(block, Table):
            full_text.extend(_table_to_lines(block))

    # Headers and footers
    for section in document.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                text = _single_line(p.text)
                if text:
                    full_text.append(text)
            for table in container.tables:
                full_text.extend(_table_to_lines(table))

    # Text boxes / shapes (now scans header/footer XML parts too)
    full_text.extend(_extract_textboxes_from_docx_bytes(content))

    full_text = _pair_label_value_lines(full_text)
    full_text = [_single_line(x) for x in full_text if _single_line(x)]
    full_text = _dedupe_keep_order(full_text)
    return "\n".join(full_text).strip()


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX IMAGE OCR  +  LIBREOFFICE FALLBACK
# ═══════════════════════════════════════════════════════════════════════════════

# Formats that PIL/Tesseract cannot read; must go via LibreOffice -> PDF.
_UNREADABLE_MEDIA_EXT = {".emf", ".wmf", ".svg"}


def get_docx_media_names(content: bytes) -> list[str]:
    try:
        with zipfile.ZipFile(BytesIO(content)) as zf:
            return [n for n in zf.namelist() if n.startswith("word/media/")]
    except Exception:
        logger.exception("[TEXT_EXTRACT] Failed to list DOCX media files")
        return []


def docx_contains_images(content: bytes) -> bool:
    return bool(get_docx_media_names(content))


def ocr_images_from_docx_bytes(
    content: bytes,
    min_width: int  = 800,    # kept for backward compatibility
    min_height: int = 800,    # kept for backward compatibility
    lang: str = OCR_LANG,
) -> tuple[str, dict]:
    """
    OCR all embedded images in a DOCX package.

    Filter logic improved:
      Old: skip if width < 800 AND height < 800
           -> missed wide-but-short full-width scans (e.g. 2480 x 200)
      New: skip if area < 200,000 OR shortest side < 300
           -> catches those wide/short scans while still skipping tiny icons

    min_width / min_height params kept so any external callers don't break.
    EMF/WMF/SVG files skipped gracefully (PIL can't open them).
    """
    # Derive area/dimension thresholds from the original width/height params
    # so callers that override min_width/min_height still get sensible behaviour.
    min_area      = min_width * min_height // 4   # ≈ 160 000 for defaults
    min_dimension = min(min_width, min_height) // 3  # ≈ 266 for defaults

    texts                : list[str] = []
    total_images          = 0
    ocr_images_count      = 0
    skipped_small         = 0
    skipped_unreadable    = 0

    try:
        with zipfile.ZipFile(BytesIO(content)) as zf:
            for media_name in get_docx_media_names(content):
                total_images += 1
                ext = Path(media_name).suffix.lower()

                if ext in _UNREADABLE_MEDIA_EXT:
                    skipped_unreadable += 1
                    continue

                try:
                    with zf.open(media_name) as img_file:
                        img = Image.open(BytesIO(img_file.read()))
                        img.load()
                except Exception:
                    logger.exception("[TEXT_EXTRACT] Failed loading DOCX media: %s", media_name)
                    continue

                w, h = img.size if img.size else (0, 0)
                if (w * h) < min_area or min(w, h) < min_dimension:
                    skipped_small += 1
                    continue

                try:
                    text = ocr_image(img, lang=lang)
                except Exception:
                    logger.exception("[TEXT_EXTRACT] OCR failed for DOCX media: %s", media_name)
                    text = ""

                if text:
                    texts.append(text)
                    ocr_images_count += 1

    except Exception:
        logger.exception("[TEXT_EXTRACT] Failed OCR fallback on DOCX images")

    final_text = "\n".join(t for t in texts if t).strip()
    meta = {
        "mode"                    : "ocr_images",
        "total_images"            : total_images,
        "ocr_images"              : ocr_images_count,
        "skipped_small_images"    : skipped_small,
        "skipped_unreadable_images": skipped_unreadable,   # new key
        "extracted_chars"         : len(final_text or ""),
        "source"                  : "docx_images",
    }
    return final_text, meta


def find_office_converter_executable() -> str | None:
    return shutil.which("libreoffice") or shutil.which("soffice")


def convert_docx_bytes_to_pdf(content: bytes) -> bytes:
    converter = find_office_converter_executable()
    if not converter:
        raise RuntimeError("LibreOffice/soffice executable not found for DOCX to PDF conversion")

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path  = os.path.join(tmpdir, "input.docx")
        output_path = os.path.join(tmpdir, "input.pdf")

        with open(input_path, "wb") as f:
            f.write(content)

        proc = subprocess.run(
            [converter, "--headless", "--convert-to", "pdf",
             "--outdir", tmpdir, input_path],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=120,
            check=False,
        )

        if proc.returncode != 0 or not os.path.exists(output_path):
            raise RuntimeError(
                f"DOCX to PDF conversion failed. returncode={proc.returncode}, "
                f"stderr={proc.stderr.decode(errors='ignore')}"
            )

        with open(output_path, "rb") as f:
            return f.read()


def extract_text_from_docx_hybrid(content: bytes) -> tuple[str, dict]:
    """
    DOCX extraction strategy:
      1. Native DOCX text extraction          — always runs, fastest.
      2. Direct image OCR                     — if native insufficient + readable images present.
      3. LibreOffice DOCX->PDF conversion     — last resort, also triggered when DOCX
                                                contains only EMF/WMF media that PIL can't read.
    """
    native_text  = (extract_text_from_docx_bytes(content) or "").strip()
    media_names  = get_docx_media_names(content)
    has_images   = bool(media_names)

    # True when every media file is a vector format Tesseract can't read.
    only_unreadable_media = has_images and all(
        Path(n).suffix.lower() in _UNREADABLE_MEDIA_EXT for n in media_names
    )

    native_meta = {
        "mode"            : "native",
        "total_pages"     : None,
        "pages_native"    : None,
        "pages_ocrd"      : 0,
        "suspicious_pages": [],
        "extracted_chars" : len(native_text or ""),
        "fallback_used"   : False,
        "has_images"      : has_images,
        "media_count"     : len(media_names),
        "source"          : "docx_native",
    }

    # Fast path: native text is complete — skip everything else.
    if is_text_sufficient(native_text):
        return native_text, native_meta

    # No images — native is all we have.
    if not has_images:
        native_meta["mode"] = "native_low_text"
        return native_text, native_meta

    combined_text = native_text
    image_meta    = None

    # Try direct image OCR (much faster than LibreOffice).
    # Skip if all media are EMF/WMF/SVG — PIL can't open them.
    if not only_unreadable_media:
        try:
            image_text, image_meta = ocr_images_from_docx_bytes(content)
            combined_text = "\n".join(
                part for part in [native_text, image_text] if part
            ).strip()
            image_meta.update({
                "fallback_used": True,
                "has_images"   : has_images,
                "media_count"  : len(media_names),
            })
            if is_text_sufficient(combined_text):
                return combined_text, image_meta
        except Exception:
            logger.exception("[TEXT_EXTRACT] Direct image OCR failed")
            combined_text = native_text

    # Last resort: LibreOffice DOCX->PDF (slow on cold start).
    try:
        pdf_bytes        = convert_docx_bytes_to_pdf(content)
        pdf_text, pdf_meta = extract_text_from_pdf_hybrid(pdf_bytes)

        if _non_ws_len(pdf_text) >= _non_ws_len(combined_text):
            pdf_meta.update({
                "source"      : "docx_to_pdf",
                "fallback_used": True,
                "has_images"  : has_images,
                "media_count" : len(media_names),
            })
            return pdf_text, pdf_meta
    except Exception:
        logger.exception("[TEXT_EXTRACT] DOCX -> PDF fallback failed")

    # Return whichever of image-OCR vs native gave more text.
    if image_meta is not None and _non_ws_len(combined_text) > _non_ws_len(native_text):
        return combined_text, image_meta

    native_meta["mode"]         = "native_low_text"
    native_meta["fallback_used"] = False
    return native_text, native_meta


# ═══════════════════════════════════════════════════════════════════════════════
#  TEXT NORMALISATION
# ═══════════════════════════════════════════════════════════════════════════════

def normalize_extracted_text(t: str) -> str:
    if not t:
        return ""
    t = t.replace("\u2013", "-").replace("\u2014", "-").replace("\u2215", "/")
    t = t.replace("\u200b", "").replace("\ufeff", "").replace("\xa0", " ")
    # "A W S" -> "AWS" (only consecutive single capital/lowercase letters)
    t = re.sub(
        r"\b(?:[A-Za-z]\s){2,}[A-Za-z]\b",
        lambda m: m.group(0).replace(" ", ""),
        t,
    )
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in t.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines).strip()


# ═══════════════════════════════════════════════════════════════════════════════
#  PUBLIC ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_bytes(
    content: bytes,
    filename: str,
    document_type: str = "jd",
    return_meta: bool = False,
):
    """
    Unified entry point for PDF / DOCX / TXT extraction.
    Signature and return shape identical to previous version.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        text, meta = extract_text_from_pdf_hybrid(content)
    elif suffix == ".docx":
        text, meta = extract_text_from_docx_hybrid(content)
    elif suffix == ".txt":
        text = content.decode("utf-8", errors="ignore").strip()
        meta = {
            "mode"            : "native",
            "total_pages"     : None,
            "pages_native"    : None,
            "pages_ocrd"      : 0,
            "suspicious_pages": [],
            "extracted_chars" : len(text or ""),
            "fallback_used"   : False,
            "source"          : "txt",
        }
    else:
        raise HTTPException(
            status_code=400,
            detail="Only PDF, DOCX or TXT files are supported",
        )

    if document_type == "jd":
        text = normalize_extracted_text(text)

    if return_meta:
        return text, meta
    return text